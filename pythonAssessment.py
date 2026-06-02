import re
from docx import Document
from collections import Counter
from pathlib import Path

#sample_text = "Lucas goes to school every day of the week. He has many subjects to go to each school day: English, art, science, mathematics, gym, and history. His mother packs a big backpack full of books and lunch for Lucas.His first class is English, and he likes that teacher very much. His English teacher says that he is a good pupil, which Lucas knows means that she thinks he is a good student."
sample_text = ""
#f = open("")
doc_path = Path("newsarticle.docx")
print(doc_path.suffix)
word_doc = Document("newsarticle.docx")
word_text = "\n".join([paragraph.text for paragraph in word_doc.paragraphs])
#print(word_text)

def count_specific_word(text, word_search):
    
    if not text or not word_search:
        return 0
    
    pattern = r'\b' + re.escape(word_search.lower()) + r'\b'
    matches = re.findall(pattern, text.lower())
    
    count = len(matches)
    print(f"The word '{word_search}' appears {count} times")
    return count

def identify_most_common_word(text):
    
    if not text:
        return None
    
    pattern = r"\b\w+\b"
    words = re.findall(pattern, text.lower())
    word_counts = Counter(words)
    
    most_common_word = word_counts.most_common(1)
    if most_common_word:
        top_word = most_common_word[0]

        word = top_word[0]

        print("Most common word:", word)
        return word
    else:
        print("No words found")
    
     
    

def calculate_average_word_length(text):
    if not text:
        return 0
    pattern = r"\b\w+\b"
    words = re.findall(pattern, text)
    
    total_length = sum(len(word) for word in words)
    
    average_length = total_length / len(words)
    print("Average length is", float(average_length))
    
    return float(average_length)

def count_paragraph(text):
    
    if not text:
        print("No paragraph")
        return 1
    
    count = 0
    paragraphs = text.paragraphs
    while count < len(paragraphs):
        count += 1
    
    print("Paragraphs:", count)
    return int(count)

def count_sentences(text):
    
    if not text:
        print("No sentences")
        return 1
    
    pattern = r"(?<=[.!?])+"
    sentences = re.split(pattern, text)
    clean_sentences = []
    for sentence in sentences:
        if sentence.strip() != "":
            clean_sentences.append(sentence)
    
    sentence_count = len(clean_sentences)
    print("Number of sentences:", sentence_count)
    return int(sentence_count)
    


identify_most_common_word(word_text)
calculate_average_word_length(word_text)
count_paragraph(word_doc)
count_sentences(sample_text)
count_specific_word(word_text,"machine")